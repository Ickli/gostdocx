'''
You can use the script as... a script. An example is in example.py and README.

Also, use it as a library! Just use init_gostdocx(**kwargs) to pass arguments
to the module (as if through command line) and then process_txt(inpath, outpath)
'''

import os
import sys
import argparse
from docx import Document
from GdocxState import GdocxState
from typing import Type, Any
import GdocxParsing
import GdocxHandler
import GdocxStyle
import GdocxCommon
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns
from docxcompose.composer import Composer

STARTUP_INPUT_DIR = "."
# Relative path is resolved against the script's path,
# as the default styles are intended to be in its directory.
PATH_DEFAULT_STYLES = "styles/default.json"
SKIP_NUMBERING = False
CONVERT_DOCX_TO_TXT = False
DOCX_TO_TXT_OUTDIR = "."

# ! You can add something here !
# Will be added to GdocxState's registered_handlers
registered_macro_handlers: list[Type[Any]] = [
    GdocxHandler.EchoHandler,
    GdocxHandler.ChdirHandler,
]

def process_with_current_handler(file, state: GdocxState):
    # this is for one-line macro
    if state.reached_macro_end:
        state.handler.finalize()
        state.reached_macro_end = False
        return

    line = file.readline()
    while(line != ""):
        new_handler = state.handle_or_get_new_handler(line)

        if new_handler is not None:
            # new_handler sees old_handler as state's current handler.
            # so it can perform some logic based on that
            old_handler = state.handler
            state.handler = new_handler

            if old_handler.NAME == GdocxState.NAME:
                state.finalize()

            state.indent += 1
            process_with_current_handler(file, state)

            state.indent -= 1
            state.handler = old_handler

            if state.reached_page_macro:
                return

        # this is for end macro.
        # page macro is one-liner, so it ends immediately and we can return something
        if state.reached_macro_end:
            state.handler.finalize()
            state.reached_macro_end = False
            return

        line = file.readline()
    return

# Copied from https://stackoverflow.com/questions/56658872/add-page-number-using-python-docx
def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
# end copied

def add_footer_with_page_number(doc: Document):
    add_page_number(doc.sections[0].footer.paragraphs[0].add_run())
    doc.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def process_txt(filepath: str, filepath_out: str):
    file = open(filepath, "r")
    doc = Document()
    GdocxStyle.use_default_styles(doc)
    docs = []

    while True:
        with GdocxState(doc, registered_macro_handlers) as state:
            state.strip_indent = GdocxParsing.STRIP_INDENT
            state.skip_empty = GdocxParsing.SKIP_EMPTY
            # here state is primary handler
            process_with_current_handler(file, state)
            to_append = state.reached_page_macro

            docs.append(doc)
            if to_append:
                docs.append(Document(state.append_filepath))
                doc = Document()
                GdocxStyle.use_default_styles(doc)
            else:
                break

    file.close()

    if not SKIP_NUMBERING:
        add_footer_with_page_number(docs[0])

    composer = Composer(docs[0])
    for doc in docs[1:]:
        composer.append(doc)
    composer.save(filepath_out)


def process_args() -> (str, str):
    inpath: str | None = None
    outpath: str | None = None
    docx_to_txt_outdir: str | None = None

    prs = argparse.ArgumentParser(prog = "GostDocx",
        description = '''
Converts .txt files into .docx
>> To see an example, run 'python3 example.py'

Конвертирует .txt файлы в .docx
>> Чтобы увидеть пример, введи 'python3 example.py' ''',
        formatter_class = argparse.RawTextHelpFormatter
    )
    prs.add_argument('-i', '--input', help="Path to source file", type=str)
    prs.add_argument('-o', '--output', help="Path to output file, or outut file name in case -d flag is provided", type=str)
    prs.add_argument('-s', '--strip-indent', help="strip indents of nested macros", action="store_true")
    prs.add_argument('-se', '--skip-empty', help="skip empty lines", action="store_true")
    prs.add_argument('-il', '--indent-length', help="Length of indent sequence", type=int)
    prs.add_argument('-ic', '--indent-char', help="Indent character", type=str)
    prs.add_argument('-n', '--skip-numbering', help="Don't put page number in footers of pages", action="store_true")
    prs.add_argument('-d', '--docx_to_txt', help="Convert .docx file .txt", action="store_true")
    prs.add_argument('-od', '--docx_to_txt_outdir', help="If -d flag is provided, specifies output dir for style and output files", type=str)
    prs.add_argument('-id', '--input_dir', help="Program moves to specified directory before processing txt's. If not specified, uses current working dir. Paths passed via -i and -o are resolved before moving", type=str)

    args = prs.parse_args()
    inpath = args.input
    outpath = args.output

    if inpath == None:
        print("ERROR: must provide path to in file .txt")
        exit(1)
    elif outpath == None:
        print("ERROR: must provide path to out file .docx")
        exit(1)

    init_gostdocx(
        indent_length = args.indent_length,
        indent_char = args.indent_char,
        input_dir = args.input_dir,
        strip_indent = args.strip_indent,
        skip_empty = args.skip_empty,
        skip_numbering = args.skip_numbering,
        docx_to_txt_outdir = args.docx_to_txt_outdir,
        docx_to_txt = args.docx_to_txt
    )

    return (inpath, outpath)

def init_gostdocx(**kwargs):
    il = kwargs.get('indent_length')
    ic = kwargs.get('indent_char')
    input_dir = kwargs.get('input_dir')

    if ic == None:
        ic = GdocxParsing.INDENT_DEFAULT_CHAR
    if il == None:
        il = GdocxParsing.INDENT_DEFAULT_LENGTH
    GdocxParsing.INDENT_STRING = ic * il
    
    if input_dir is not None:
        global STARTUP_INPUT_DIR
        STARTUP_INPUT_DIR = GdocxCommon.AbsPath(input_dir)

    GdocxParsing.STRIP_INDENT = kwargs.get('strip_indent')
    GdocxParsing.SKIP_EMPTY = kwargs.get('skip_empty')
    SKIP_NUMBERING = kwargs.get('skip_numbering')

    CONVERT_DOCX_TO_TXT = kwargs.get('docx_to_txt')
    if CONVERT_DOCX_TO_TXT:
        od = kwargs.get('docx_to_txt_outdir')
        if od is None:
            print("ERROR: Must provide -od flag when -d flag is provided")
            exit()
        global DOCX_TO_TXT_OUTDIR 
        DOCX_TO_TXT_OUTDIR = od


if __name__ == "__main__":
    inpath, outpath = process_args()

    # resolve abs paths before switching to STARTUP_INPUT_DIR
    inpath = GdocxCommon.AbsPath(inpath)
    outpath = GdocxCommon.AbsPath(outpath)

    # init default styles, which is in a sibling file to the script
    absScriptPath = os.path.dirname(os.path.realpath(__file__))
    absDefaultStylesPath = os.path.join(absScriptPath, PATH_DEFAULT_STYLES)
    GdocxStyle.init_default_styles(absDefaultStylesPath)

    # if was specified in args, change to it.
    # otherwise, just stays in the directory of the caller
    os.chdir(STARTUP_INPUT_DIR)

    if not CONVERT_DOCX_TO_TXT:
        process_txt(inpath, outpath)
        print(f"\'{outpath}\' created")
    else:
        print("DOCX_TO_TXT feature is not yet implemented")
        exit(1)

        import GdocxToTxt
        outname = outpath
        GdocxToTxt.docx_to_txt(inpath, outname, DOCX_TO_TXT_OUTDIR)
        print(f"{docx_to_txt_outdir} dir, '{docx_to_txt_outdir}/{outname}', '{docx_to_txt_outdir}/styles.json' created")
