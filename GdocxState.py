from typing import Type, Any
from docx import Document
from docx.styles.style import ParagraphStyle, CharacterStyle
from docx.text.paragraph import Paragraph
from docx.text.run import Run
import GdocxHandler
import GdocxParsing
import GdocxStyle

default_handlers: list[Type[Any]] = [
        GdocxHandler.OrderedListHandler,
        GdocxHandler.OrderedListItemHandler,
        GdocxHandler.UnorderedListHandler,
        GdocxHandler.UnorderedListItemHandler,
        GdocxHandler.LoadStyleHandler,
        GdocxHandler.ParStyleHandler,
        GdocxHandler.ImageHandler,
        GdocxHandler.ImageCaptionHandler,
        GdocxHandler.PageBreakHandler,
        GdocxHandler.TableHandler,
        GdocxHandler.TableCellHandler,
        GdocxHandler.AppendPageHandler,
        GdocxHandler.JsonReaderHandler,
        GdocxHandler.JsonFieldHandler,
        GdocxHandler.RunStyleHandler,
        GdocxHandler.ImageNumberAsRunHandler,
        GdocxHandler.NextImageNumberAsRunHandler,
        GdocxHandler.SpaceHandler,
]

# Document passed to ctor must outlive GdocxState.
class GdocxState:
    NAME = "gdocx-state"
    STYLE = "paragraph"

    def __init__(self, 
        doc: Document, 
        handlers: list[Type[Any]]
    ):
        self.doc = doc
        # almost all handlers refer to state.receiver and not state.doc
        self.receiver = GdocxStateReceiver(self)
        self.paragraph_lines = []
        self.current_style = doc.styles['Normal']
        self.registered_handlers = get_handler_dict(default_handlers + handlers)
        self.reached_macro_end = False
        self.handler = self
        self.indent = 0
        self.strip_indent = False
        self.skip_empty = False
        self.line_number = 0

        self.reached_page_macro = False
        self.append_filepath = ""

    # Returns new handler, if macro is encountered;
    # otherwise, returns None
    def handle_or_get_new_handler(self,
        line: str
    ) -> object | None:
        self.line_number += 1
        indent = self.indent if self.strip_indent else 0

        rawline, info = GdocxParsing.parse_line(line, indent)

        if info.is_empty and self.skip_empty:
            return None
        if info.type == GdocxParsing.INFO_TYPE_MACRO:
            return self.process_macro_line(rawline, info)
        elif info.type != GdocxParsing.INFO_TYPE_COMMENT:
            self.handler.process_line(rawline, info)

        return None

    def process_line(self, line: str, info: GdocxParsing.LineInfo):
        # GdocxParsing.INFO_TYPE_MACRO is handled in caller 'handle_or_get_new_handler'
        self.paragraph_lines.append(info.line_stripped)
        self.finalize()

    # Returns new handler or None
    def process_macro_line(self, 
        line: str, 
        info: GdocxParsing.LineInfo
    ) -> object | None:
        macro_type = GdocxParsing.get_macro_type(info.line_stripped)
        new_handler = None

        if macro_type == GdocxParsing.MACRO_TYPE_START or macro_type == GdocxParsing.MACRO_TYPE_ONE_LINE:
            args = GdocxParsing.parse_macro_args(info.line_stripped)
            if len(args) == 0:
                raise Exception("Empty macro")
            macro_name = args[0]
            if macro_name not in self.registered_handlers:
                raise Exception("Couldn't find macro: %s" % macro_name)
            else:
                new_handler = self.registered_handlers[macro_name](self, args[1:])

        self.reached_macro_end = (macro_type == GdocxParsing.MACRO_TYPE_END or macro_type == GdocxParsing.MACRO_TYPE_ONE_LINE)

        return new_handler
    
    def process_header(self, line: str, info: GdocxParsing.LineInfo):
        self.doc.add_heading(GdocxParsing.get_header_string(line), 0)

    def finalize(self):
        if len(self.paragraph_lines) != 0:
            par_content = '\n'.join(self.paragraph_lines)
            self.doc.add_paragraph(par_content, style = self.STYLE)
            self.paragraph_lines = []

    def __enter__(self):
        return self
    
    def __exit__(self, exc_type, exc_value, traceback):
        self.finalize()


def get_handler_dict(handlers: list[Type[Any]]):
    handler_dict = {}
    for handler in handlers:
        handler_dict[handler.NAME] = handler
    return handler_dict


class GdocxStateReceiver:
    NAME = "GdocxStateReceiver"

    def __init__(self, state: GdocxState):
        self.state = state

    def add_paragraph(self, text: str = '', style: str | ParagraphStyle | None = None) -> Paragraph:
        return self.state.doc.add_paragraph(text, style)

    def add_run(self, text: str = '', style: str | CharacterStyle | None = None) -> Run:
        return self.get_paragraphs()[-1].add_run(text, style)

    def get_paragraphs(self):
        return self.state.doc.paragraphs
