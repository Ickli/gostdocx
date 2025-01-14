from docx2python import docx2python
from docx import Document
from docx.text.paragraph import Paragraph
from docx2python import docx_output
import GdocxHandler

FIRST_PROMPT = '''You'll be presented with contents of .docx file.
The contents are separated into tables. You'll see first tables
to decide if you want to start parsing with them. If not, you can
skip these tables (that's done so because some things in docx file
are not quite representable in txt format).

Press 'y' and 'Enter' to start parsing from this table,

If you want to skip current table, just press Enter or anything else.
Press 'Enter' to proceed.
'''

YES_STR = "y"
YES_NO_PROMPT = f"Press '{YES_STR}' to start parsing, press 'Enter' to skip"

class PythonDocxInfo:
    def __init__(self):
        # Set before parsing
        self.doc = None
        self.first_table_index = 0 # it's for true table
        self.first_par_index = 0
        self.first_image_index = 0

        # Used during parsing
        self.run_index = 0
        self.par_index = 0

class Docx2PythonInfo:
    def __init__(self):
        # Set before parsing
        self.doc = None
        self.first_table_index = 0

        # Used during parsing

def append_page_break(instr: str) -> str:
    instr += f"({GdocxHander.PageBreakHandler.NAME})\n"

def is_true_table(body_runs_elem: list[list[list[list[str]]]]) -> bool:
    return len(body_runs_elem) > 1

IS_IMAGE_NOT = 0
IS_IMAGE_SINGLE = 1
IS_IMAGE_DOUBLE = 2

# checks for boundaries
# if out of bounds, returns IS_IMAGE_NOT
def is_image(d2p_par: list[str], d2p_run_index) -> int:
    # TODO
    return IS_IMAGE_NOT


MET_NONE = 0
MET_IMAGE = 1
MET_PAGE_BREAK = 2
# returns count of runs to freely fetch into a paragraph
# and a flag marking condition of first encountered run not suitable
# for this paragraph
def fetch_pd_runs(pd_par: Paragraph, start_run_index: int) -> (int, int):
    

def pd_runs_to_str(pd_par: Paragraph, start_run_index: int) -> str:
    # TODO
    pass

def is_empty(body_runs: list[list[list[list[str]]]]) -> bool:
    return len(body_runs[0][0][0]) == 0

# assumes prev_body_par doesn't end with image tags
def has_page_break(prev_body_par: list[str], python_docx_run: str) -> bool:
    return len(prev_body_par) != 0 and prev_body_par[0].endswith('\n') and python_docx_run == ''

def docx_to_txt_prompted(filepath: str, outputpath: str):
    print(FIRST_PROMPT)
    input()

    docx2pythoni = Docx2PythonInfo()
    python_docxi = PythonDocxInfo()

    docx2pythoni.doc = docx2python(filepath)
    python_docxi.doc = Document(filepath)

    body_runs = docx2pythoni.doc.body_runs

    for cur_table_index in range(len(body_runs)):
        print(YES_NO_PROMPT)
        inputstr = input()

        if inputstr == YES_STR:
            first_table_index = cur_table_index
            break

        cur_table = body_runs[cur_table_index];
        if is_true_table(cur_table):
            python_docxi.first_table_index += 1
            continue

        par_count = len(cur_table[0][0]) # 0'th row, 0'th cell
        python_docxi.first_par_index += par_count

    docx_to_txt(python_docxi, docx2pythoni, outputpath)
    
def docx_to_txt(python_docxi: PythonDocxInfo, docx2pythoni: Docx2PythonInfo, outputpath: str):
    print(docx_to_text(python_docxi, docx2pythoni)
    return

def docx_to_text(python_docxi: PythonDocxInfo, docx2pythoni: Docx2PythonInfo) -> str:
    d2p_tables = docx2pythoni.doc.body_runs[docx2pythoni.first_table_index:]
    d2p_tables_len = len(d2p_tables)

    pd_pars = python_docxi.doc.paragraphs[python_docxi.first_par_index:]
    pd_tables = python_docxi.doc.tables[python_docxi.first_par_index:]
    pd_pars_len = len(pd_pars)
    pd_tables_len = len(pd_tables)
    
    pd_par_index = 0
    d2p_par_index = 0
    d2p_table_index = 0

    outstr = ""

    while d2p_table_index < d2p_tables_len:
        d2p_table = d2p_tables[d2p_table_index]

        if is_true_table(d2p_table):
            # TODO: process table
            continue

        d2p_pars = d2p_table[0][0]
        d2p_pars_len = len(d2p_pars)
        d2p_par_index = 0
        pd_run_index = 0

        while d2p_par_index < d2p_pars_len:
            d2p_par = d2p_pars[d2p_par_index]
            d2p_par_len = len(d2p_par)
            if d2p_par_len == 0:
                pd_par_index += 1
                d2p_par_index += 1
                continue

            d2p_run_index = 0

            while d2p_run_index < d2p_par_len:
                pd_run_index = 0
                d2p_run = d2p_par[d2p_run_index]

                isimage = IS_IMAGE_NOT
                if d2p_run.endswith('\n'):
                    isnextimage = is_image(d2p_par, d2p_run_index + 1)

                    if not isnextimage:
                        outstr = append_page_break(outstr)
                        pd_run_index += 1
                    else: # TODO: image logic
                        pass
                elif '\n' in d2p_run:
                    # TODO: change main logic to use last_condition
                    fetched_count, last_condition = fetch_pd_runs(
                        pd_pars[pd_par_index], pd_run_index)
                    outstr += pd_runs_to_str(pd_pars[pd_par_index], pd_run_index)
                    pd_run_index += fetched_count
                elif not (isimage = is_image(d2p_par, d2p_run_index)):
                    fetched_count, last_condition = fetch_pd_runs(
                        pd_pars[pd_par_index], pd_run_index)
                    outstr += pd_runs_to_str(pd_pars[pd_par_index], pd_run_index)
                    pd_run_index += fetched_count
                else:
                    # TODO: image logic
                    pass

            pd_par_index += 1
            d2p_par_index += 1
