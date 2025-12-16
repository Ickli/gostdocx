import GdocxParsing
import GdocxStyle
import GdocxCommon
import os.path
import json
from docx.shared import Cm
from docx.table import _Cell
from docx.styles.style import ParagraphStyle, CharacterStyle
from docx.text.paragraph import Paragraph
from docx.text.run import Run

'''
handler trait:
    NAME: str

    Called on every occurrence of custom macro in file
    __init__(self, state: GdocxState, macro_args: list[str]);

    Called on every line inside custom macro
    process_line(self, line: str, info: GdocxParsing.LineInfo);

    Called when reached end of scope of custom macro
    finalize(self);
'''

'''
Almost all handlers refer to state.receiver and not state.doc to
append paragraphs. It serves as a proxy to state.doc.
Handlers may change state.receiver

receiver trait:
    NAME: str

    add_paragraph(self, text: str = '', style: str | ParagraphStyle | None = None) -> Paragraph;

    add_run(self, text: str = '', style: str | CharacterStyle | None = None) -> Run;

    get_paragraphs(self) -> list[docx.Paragraph]
'''

class EchoHandler:
    NAME = "echo"
    Prefix = "echo: "

    def __init__(self, state: 'GdocxState', macro_args: list[str]):
        print(self.Prefix, *macro_args)
        pass

    def process_line(self, line: str, info: GdocxParsing.LineInfo):
        raise Exception(f"You must not place content inside {self.NAME}")

    def finalize(self):
        pass

class ChdirHandler:
    NAME = "chdir"

    def __init__(self, state: 'GdocxState', macro_args: list[str]):
        self.ddir = macro_args[0]
        pass

    def process_line(self, line: str, info: GdocxParsing.LineInfo):
        raise Exception(f"You must not place content inside {self.NAME}")

    def finalize(self):
        import os
        print(f"Changing dir to {self.ddir}");
        os.chdir(self.ddir)
        pass

class PageBreakHandler:
    NAME = "page-break"

    def __init__(self, state: 'GdocxState', macro_args: list[str]):
        self.state = state
        pass

    def process_line(self, line: str, info: GdocxParsing.LineInfo):
        raise Exception(f"You must not place content inside {self.NAME}")

    def finalize(self):
        self.state.doc.add_page_break()

class UnorderedListHandler:
    NAME = "unordered-list"

    def __init__(self, state: 'GdocxState', macro_args: list[str]):
        pass

    def process_line(self, line: str, info: GdocxParsing.LineInfo):
        raise Exception(f"You musn't place contents inside {self.NAME} macro")

    def finalize(self):
        pass

# Converts paragraphs to unordered list items
class UnorderedListItemHandler:
    NAME = "unordered-list-item"
    STYLE = "list"

    def __init__(self, state: 'GdocxState', macro_args: list[str]):
        if state.handler.NAME != UnorderedListHandler.NAME:
            raise Exception(f"{self.NAME} macro must be inside {UnorderedListHandler.NAME} macro")
        self.state = state
        self.cur_paragraph_lines = []
        self.is_first_line_processed = False
        state.handler.finalize()

    def process_line(self, line: str, info: GdocxParsing.LineInfo):
        if not self.is_first_line_processed:
            line_stripped = GdocxStyle.Style.UNORDERED_LIST_PREFIX + info.line_stripped
            self.is_first_line_processed = True
        self.cur_paragraph_lines.append(line_stripped)

    def finalize(self):
        par_content = '\n'.join(self.cur_paragraph_lines)
        par = self.state.receiver.add_paragraph(par_content, style = self.STYLE)

# Applies style to macro contents, treating it as a paragraph
class ParStyleHandler:
    NAME = "paragraph-styled"

    def __init__(self, state: 'GdocxState', macro_args: list[str]):
        if len(macro_args) == 0:
            raise Exception("Style macro must contain style name")
        self.state = state
        self.style_name = macro_args[0]
        self.cur_paragraph_lines = []

    def process_line(self, line: str, info: GdocxParsing.LineInfo):
        self.cur_paragraph_lines.append(info.line_stripped)

    def finalize(self):
        par_content = '\n'.join(self.cur_paragraph_lines)
        par = self.state.receiver.add_paragraph(par_content)
        par.style = self.state.doc.styles[self.style_name]

class LoadStyleHandler:
    NAME = "load-style"

    def __init__(self, state: 'GdocxState', macro_args: list[str]):
        if len(macro_args) == 0:
            raise Exception("ParseStyleDirective macro must contain path to file")
        to_override = False
        if len(macro_args) > 1:
            to_override = bool(macro_args[1])
        GdocxStyle.use_styles_from_file(macro_args[0], state.doc, to_override)

    def process_line(self, line: str, info: GdocxParsing.LineInfo):
        raise Exception("You must not place content inside ParseStyleDirective")

    def finalize(self):
        pass

class OrderedListHandler:
    NAME = "ordered-list"

    def __init__(self, state: 'GdocxState', macro_args: list[str]):
        self.free_item_number = 1
        state.handler.finalize()

    def process_line(self, line: str, info: GdocxParsing.LineInfo):
        raise Exception(f"You must provide contents of {self.NAME} within macro {OrderedListItemHandler.NAME}")

    def finalize(self):
        pass


class OrderedListItemHandler:
    INFIX = ". "
    NAME = "ordered-list-item"
    STYLE = "list"

    def __init__(self, state: 'GdocxState', macro_args: list[str]):
        if state.handler.NAME != OrderedListHandler.NAME:
            raise Exception(f"{self.NAME} macro must be inside {OrderedListHandler.NAME} macro")

        self.state = state
        self.cur_paragraph_lines = []
        self.is_first_line_processed = False

        self.item_number = state.handler.free_item_number
        state.handler.free_item_number += 1

    def process_line(self, line: str, info: GdocxParsing.LineInfo):
        if not self.is_first_line_processed:
            line_stripped = str(self.item_number) + self.INFIX + info.line_stripped
            self.is_first_line_processed = True
        self.cur_paragraph_lines.append(line_stripped)

    def finalize(self):
        par_content = '\n'.join(self.cur_paragraph_lines)
        self.state.receiver.add_paragraph(par_content, style = self.STYLE)

class ImageHandler:
    NAME = "image"
    STYLE = "image-paragraph"

    def __init__(self, state: 'GdocxState', macro_args: list[str]):
        if len(macro_args) == 0:
            raise Exception(f"Macro {self.NAME} called without arguments")
        self.state = state
        self.width = None
        self.height = None
        
        argslen = len(macro_args)
        if argslen > 1:
            widthstr = macro_args[1]
            if widthstr != "None":
                self.width = Cm(float(widthstr))

        if argslen > 2:
            heightstr = macro_args[2]
            if heightstr != "None":
                self.height = Cm(float(heightstr))

        self.path = macro_args[0]

    def process_line(self, line: str, info: GdocxParsing.LineInfo):
        raise Exception(f"You must not put contents into {self.NAME} macro")

    def finalize(self):
        par = self.state.receiver.add_paragraph(None, style = self.STYLE)
        run = par.add_run()
        run.add_picture(self.path, self.width, self.height)

class ImageCaptionHandler:
    # Because GOST wants us to minimize distance between image and its caption,
    # we insert the latter to the paragraph of the first. Thus, it's as if user
    # presses SHIFT+ENTER and then types caption by hand
    STICK_TO_PREV_PARAGRAPH = True
    NAME = "image-caption"
    STYLE = "image-caption"
    ItemFreeNumber = 1

    def __init__(self, state: 'GdocxState', macro_args: list[str]):
        self.state = state
        self.paragraph_lines = []
        self.is_first_line_processed = False

        self.item_number = ImageCaptionHandler.ItemFreeNumber
        ImageCaptionHandler.ItemFreeNumber += 1

    def process_line(self, line: str, info: GdocxParsing.LineInfo):
        line_stripped = info.line_stripped
        if not self.is_first_line_processed:
            self.is_first_line_processed = True
            line_stripped = GdocxStyle.Style.IMAGE_CAPTION_PREFIX + str(self.item_number) + GdocxStyle.Style.IMAGE_CAPTION_INFIX + info.line_stripped

        self.paragraph_lines.append(line_stripped)

    def finalize(self):
        content = ' '.join(self.paragraph_lines)
        if not self.STICK_TO_PREV_PARAGRAPH:
            self.state.receiver.add_paragraph(content, style = GdocxStyle.Style.IMAGE_CAPTION)
        else:
            content = '\n' + content
            self.state.receiver.get_paragraphs()[-1].add_run(content)

class TableCellReceiver:
    NAME = "TableReceiver"

    def __init__(self, cell: _Cell):
        self.cell = cell
        self.first_par_added = False

    def add_paragraph(self, text: str = '', style: str | ParagraphStyle | None = None) -> Paragraph:
        if not self.first_par_added:
            self.cell.paragraphs[0].text = text
            self.cell.paragraphs[0].style = style
            self.first_par_added = True
            return self.cell.paragraphs[0]
        return self.cell.add_paragraph(text, style)

    def add_run(self, text: str = '', style: str | CharacterStyle | None = None) -> Run:
        if not self.first_par_added:
            run = self.cell.paragraphs[0].add_run(text, style)
            self.first_par_added = True
            return run
        else:
            return self.cell.paragraphs[-1].add_run(text, style)

    def get_paragraphs(self):
        return self.cell.paragraphs


class TableHandler:
    NAME = "table"

    def __init__(self, state: 'GdocxState', macro_args: list[str]):
        if len(macro_args) < 2:
            raise Exception(f"{self.NAME} must have at least 2 args")
        self.rows = int(macro_args[0])
        self.cols = int(macro_args[1])

        self.state = state
        self.table = self.state.doc.add_table(rows = self.rows, cols = self.cols)

    def process_line(self, line: str, info: GdocxParsing.LineInfo):
        raise Exception(f"You must not place content inside {self.NAME}")

    def finalize(self):
        pass

class TableCellHandler:
    NAME = "table-cell"
    STYLE = "paragraph"

    def __init__(self, state: 'GdocxState', macro_args: list[str]):
        if len(macro_args) < 2:
            raise Exception(f"{self.NAME} must have at least 2 args")
        if state.handler.NAME != TableHandler.NAME:
            raise Exception(f"{self.NAME} must be placed inside {TableHandler.NAME}")
        self.table = state.handler
        self.paragraph_lines = []

        self.state = state
        row = int(macro_args[0])
        col = int(macro_args[1])

        rowindex = row
        colindex = col

        if rowindex >= self.table.rows or rowindex < 0:
            raise Exception(f"{self.NAME} row (= {row}) number must be between 1 and {self.table.rows - 1} inclusive")
        if colindex >= self.table.cols or colindex < 0:
            raise Exception(f"{self.NAME} col (= {col}) number must be between 1 and {self.table.cols - 1} inclusive")

        self.prev_receiver = state.receiver

        cell = self.table.table.rows[rowindex].cells[colindex]
        state.receiver = TableCellReceiver(cell)

    def process_line(self, line: str, info: GdocxParsing.LineInfo):
        self.paragraph_lines.append(info.line_stripped)

    def finalize(self):
        if len(self.paragraph_lines) != 0:
            par_content = '\n'.join(self.paragraph_lines)
            self.state.receiver.add_paragraph(par_content, style = self.STYLE)
        self.state.receiver = self.prev_receiver

class AppendPageHandler:
    NAME = "doc"

    def __init__(self, state: 'GdocxState', macro_args: list[str]):
        if len(macro_args) == 0:
            raise Exception(f"{self.NAME} macro needs at least 1 argument")

        path = GdocxCommon.AbsPath(macro_args[0])
        if not os.path.isfile(path):
            raise Exception(f"{path} is not a file. You must pass a file path to {self.NAME} macro")

        state.reached_page_macro = True
        state.append_filepath = path
        pass

    def process_line(self, line: str, info: GdocxParsing.LineInfo):
        raise Exception(f"You must not place content inside {self.NAME}")

    def finalize(self):
        pass

class RunStyleHandler:
    NAME = "run-styled"

    def __init__(self, state: 'GdocxState', macro_args: list[str]):
        self.state = state
        if len(macro_args) > 0:
            self.style = state.doc.styles[macro_args[0]]
        else:
            self.style = None
        self.run_lines = []

    def process_line(self, line: str, info: GdocxParsing.LineInfo):
        self.run_lines.append(info.line_stripped)

    def finalize(self):
        run_content = '\n'.join(self.run_lines)
        self.state.receiver.add_run(run_content, self.style)

class JsonReaderHandler:
    NAME = "json-reader"

    def __init__(self, state: 'GdocxState', macro_args: list[str]):
        if len(macro_args) == 0:
            raise Exception(f"{self.NAME} macro needs at least 1 argument")

        self.state = state
        self.jsonname = macro_args[0]
        
        jsonfile = open(self.jsonname, 'r')
        self.json = json.loads(jsonfile.read())
        self.prev_receiver = self.state.receiver
        self.state.receiver = JsonReaderReceiver(self)

    def process_line(self, line: str, info: GdocxParsing.LineInfo):
        raise Exception(f"{self.NAME} does not accept free-standing text inside")

    def finalize(self):
        self.state.receiver = self.prev_receiver

    def get_json_field(self, fieldname):
        return self.json[fieldname]


# This class is purely for restraining json-field, so that it knows
# to which receiver it is going to send text
class JsonReaderReceiver:
    NAME = "JsonReaderReceiver"

    def __init__(self, jsonhandler: 'JsonReaderHandler'):
        self.jsonhandler = jsonhandler
        pass

    def add_paragraph(self, text: str = '', style: str | ParagraphStyle | None = None) -> Paragraph:
        return self.jsonhandler.prev_receiver.add_paragraph(text, style)

    def add_run(self, text: str = '', style: str | CharacterStyle | None = None) -> Run:
        return self.jsonhandler.prev_receiver.add_run(text, style)

    def get_paragraphs(self):
        return self.jsonhandler.prev_receiver.get_paragraphs()


class JsonFieldHandler:
    NAME = "json-field"

    def __init__(self, state: 'GdocxState', macro_args: list[str]):
        if len(macro_args) == 0:
            raise Exception(f"{self.NAME} macro needs at least 1 argument")
        if not isinstance(state.receiver, JsonReaderReceiver):
            raise Exception("{self.NAME} can only be used inside {JsonReaderHandler.NAME}")

        self.state = state
        self.fieldname = macro_args[0]

    def process_line(self, line: str, info: GdocxParsing.LineInfo):
        raise Exception(f"{self.NAME} does not accept free-standing text inside")

    def finalize(self):
        recv = self.state.receiver

        if len(recv.get_paragraphs()) == 0:
            raise Exception(f"Before {self.NAME} insert at least one {ParStyleHandler.NAME} so that it's possible to attach the field's value to it")

        recv.add_run(
            str(recv.jsonhandler.get_json_field(self.fieldname)))


class NextImageNumberAsRunHandler:
    NAME = "next-image-number-as-run"

    def __init__(self, state: 'GdocxState', macro_args: list[str]):
        self.state = state
        pass

    def process_line(self, line: str, info: GdocxParsing.LineInfo):
        raise Exception(f"You must not place content inside {self.NAME}")

    def finalize(self):
        self.state.receiver.add_run(str(ImageCaptionHandler.ItemFreeNumber))
        pass

class ImageNumberAsRunHandler:
    NAME = "image-number-as-run"

    def __init__(self, state: 'GdocxState', macro_args: list[str]):
        self.state = state
        pass

    def process_line(self, line: str, info: GdocxParsing.LineInfo):
        raise Exception(f"You must not place content inside {self.NAME}")

    def finalize(self):
        self.state.receiver.add_run(str(ImageCaptionHandler.ItemFreeNumber - 1))
        pass

class SpaceHandler:
    NAME = "space"

    def __init__(self, state: 'GdocxState', macro_args: list[str]):
        self.state = state
        if len(macro_args) > 0:
            self.count = int(macro_args[0])
        else:
            self.count = 1

    def process_line(self, line: str, info: GdocxParsing.LineInfo):
        raise Exception(f"You must not place content inside {self.NAME}")

    def finalize(self):
        run_content = ' ' * self.count
        self.state.receiver.add_run(run_content)

# Dispatches based on current macro.
# Accepts only args in format [(True macro_name|label)|(False macro_name|label ...)]
class NumberedHandler:
    NAME = "numbered"

    def __init__(self, state: 'GdocxState', macro_args: list[str]):
        self.state = state
        self.is_in_erasing_state = False
        self.in_macro_names = []

        if len(macro_args) > 1:
            if macro_args[0] == "True":
                self.is_in_erasing_state = True
                if len(macro_args) != 2:
                    raise ValueError(f"{self.NAME} if 1'st arg == True, "
                        "the 2nd (macro_name) must be provided. Do not provide"
                        " any more args.")
                self.in_macro_names.append(macro_args[1])

            elif macro_args[0] == "False":
                self.in_macro_names = macro_args[1:]

        self.prev_receiver = self.state.receiver
        self.state.receiver = NumberedReceiver(self)

    def process_line(self, line: str, info: GdocxParsing.LineInfo):
        raise Exception(f"You must not place content inside {self.NAME}")

    def finalize(self):
        if self.is_in_erasing_state:
            self.state.receiver.erase_macro_name(self.in_macro_names[0])
            self.state.receiver = self.prev_receiver
            return

        # account these macro names to nullify previous 'work' if it doesn't
        # uphold later in NumberedReceiver
        # self.state.receiver.PREV_IN_MACRO_NAMES = self.in_macro_names
        self.state.receiver = self.prev_receiver

class NumberedReceiver:
    NAME = "NumberedReceiver"
    START_NUMBER = 1

    NUMBER_DICTS = []
    PREV_IN_MACRO_NAMES = []

    def __init__(self, numbered_handler: 'NumberedHandler'):
        self.has_run = False
        self.numbered_handler = numbered_handler
        self.in_macro_names = self.numbered_handler.in_macro_names

        first_unmatched_index = None

        if len(self.in_macro_names) == 0 and len(NumberedReceiver.NUMBER_DICTS) == 0:
            self.set_empty_dicts_at(0, 1)

        for i in range(len(self.in_macro_names)):
            if i == len(NumberedReceiver.PREV_IN_MACRO_NAMES):
                first_unmatched_index = i
                print('overflow', i, NumberedReceiver.PREV_IN_MACRO_NAMES)
                break
            print('eq?', self.in_macro_names[i], NumberedReceiver.PREV_IN_MACRO_NAMES[i])
            if self.in_macro_names[i] != NumberedReceiver.PREV_IN_MACRO_NAMES[i]:
                first_unmatched_index = i
                break

        if first_unmatched_index is None:
            return

        new_number_dicts_len = max(
            len(NumberedReceiver.NUMBER_DICTS), len(self.in_macro_names))
        self.set_empty_dicts_at(first_unmatched_index, new_number_dicts_len)


    def add_paragraph(self, text: str = '', style: str | ParagraphStyle | None = None) -> Paragraph:
        if self.has_run:
            raise ValueError(f"{self.NAME} is supporting only 1 macro inside")

        self.has_run = True
        text = self.dispatch_on_name_and_transform(text)
        return self.numbered_handler.prev_receiver.add_paragraph(text, style)

    def add_run(self, text: str = '', style: str | CharacterStyle | None = None) -> Run:
        return self.numbered_handler.prev_receiver.add_run(text, style)

    def get_paragraphs(self):
        return self.numbered_handler.prev_receiver.get_paragraphs()
    
    def dispatch_on_name_and_transform(self, text: str):
        if self.numbered_handler.is_in_erasing_state:
            raise ValueError(f"{self.numbered_handler.NAME} is in erasing state."
                "It means, that it is created just for zero-ing out a number"
                " for some macro-name|label."
                " You must not put contents in such a macro")

        prefix = None
        print('start dispatch', self.in_macro_names)
        if len(self.in_macro_names) == 0:
            print("aaa")
            curhandler = self.numbered_handler.state.current_macro_name
            prefix = NumberedReceiver.construct_prefix([curhandler])
            print('assigning', curhandler)
            NumberedReceiver.PREV_IN_MACRO_NAMES = [curhandler]
        else:
            prefix = NumberedReceiver.construct_prefix(self.in_macro_names)
            NumberedReceiver.PREV_IN_MACRO_NAMES = self.in_macro_names

        return prefix + " " + text


    def erase_macro_name(self, name):
        for dic in NumberedReceiver.NUMBER_DICTS:
            dic[name] = self.START_NUMBER

    def set_empty_dicts_at(self, start, end):
        print("setting empty at", start, end)
        for i in range(start, end):
            if i == len(NumberedReceiver.NUMBER_DICTS):
                NumberedReceiver.NUMBER_DICTS.append({})
            else:
                NumberedReceiver.NUMBER_DICTS[i] = {}

    @classmethod
    def construct_prefix(cls, macro_names):
        assert(len(macro_names) > 0)
        assert(len(macro_names) <= len(cls.NUMBER_DICTS))

        res = ""
        for i in range(len(macro_names) - 1):
            macro_name = macro_names[i]
            dic = cls.NUMBER_DICTS[i]
            num = dic.get(macro_name)
            if num is None:
                num = cls.START_NUMBER
                dic[macro_name] = cls.START_NUMBER
            num -= 1
            res += str(num)
            res += "."

        macro_name = macro_names[-1]
        dic = cls.NUMBER_DICTS[len(macro_names) - 1]
        num = dic.get(macro_name)
        if num is None:
            num = dic[macro_name] = cls.START_NUMBER
        dic[macro_name] += 1

        res += str(num)

        return res
